import streamlit as st

# 将这行移到所有 st 命令之前
st.set_page_config(page_title="复新Vis-数据新闻多智能体工作流", layout="wide")

import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from openai import OpenAI
import os
from dotenv import load_dotenv
import json
import logging
from typing import Dict, Tuple, Optional, List
import functools
import time
import re  # Import the regular expression module
import requests
import asyncio
from crawl4ai import AsyncWebCrawler
from bs4 import BeautifulSoup
from pyecharts import options as opts
from pyecharts.charts import Bar, Pie, Line, Scatter, HeatMap, Tree, Sunburst
from pyecharts.charts import TreeMap, Boxplot
from pyecharts.globals import ThemeType
from streamlit_echarts import st_pyecharts
import numpy as np  # 确保导入numpy用于直方图计算
from io import BytesIO
from docx import Document
from docx.shared import Inches
import base64

def search_with_tavily(query):
    """使用 Tavily API 搜索相关信息"""
    try:
        TAVILY_API_KEY = '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'
        TAVILY_API_URL = 'https://api.tavily.com/search'
        
        data = {
            "api_key": TAVILY_API_KEY,
            "query": query,
            "search_depth": "basic",
            "max_results": 5,
            "language": "zh"
        }
        
        response = requests.post(TAVILY_API_URL, json=data)
        response.raise_for_status()
        result = response.json()
        
        # 提取最多3个搜索结果
        if 'results' in result:
            return [
                {
                    'title': item.get('title', ''),
                    'content': item.get('content', '')[:500],
                    'url': item.get('url', '')
                }
                for item in result['results'][:3]
            ]
        return []
    except Exception as e:
        logger.error(f"Tavily API 错误: {str(e)}")
        return []

# 选题确定阶段
def topic_selection_phase():
    st.header("第一步：数据新闻选题确定")
    
    # 初始化session state变量
    if 'topic_conversation' not in st.session_state:
        st.session_state.topic_conversation = []
    
    if 'suggested_topics' not in st.session_state:
        st.session_state.suggested_topics = []
    
    if 'selected_topic' not in st.session_state:
        st.session_state.selected_topic = None
    
    if 'skip_topic_selection' not in st.session_state:
        st.session_state.skip_topic_selection = False
        
    # 如果已经选择了选题，显示它并进入下一阶段
    if st.session_state.selected_topic:
        st.success(f"已选择的选题：{st.session_state.selected_topic}")
        return True
    
    # 显示聊天历史
    for message in st.session_state.topic_conversation:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
    
    # 用户输入选题方向
    topic_description = st.chat_input(
        "请描述你感兴趣的数据新闻选题方向...",
        key="topic_input"
    )
    
    # 当用户提交选题描述
    if topic_description:
        # 添加用户消息到聊天历史
        st.session_state.topic_conversation.append({"role": "user", "content": topic_description})
        
        # 显示用户消息
        with st.chat_message("user"):
            st.markdown(topic_description)
        
        # 先尝试获取相关搜索结果
        search_results = search_with_tavily(topic_description)
        search_context = ""
        
        if search_results:
            search_context = "基于以下最新资讯:\n" + "\n\n".join([
                f"标题: {result['title']}\n内容: {result['content']}\n来源: {result['url']}"
                for result in search_results
            ])
        
        # 构建提示
        system_prompt = """你是一位专业有趣的数据新闻编辑，擅长帮助记者确定有价值的数据新闻选题。
        请根据用户的选题方向，生成三个明确具体的数据新闻选题建议。每个选题必须:
        1. 具有新闻价值和数据驱动特性
        2. 明确定义了研究问题和可能的数据来源
        3. 有潜在的社会影响或公众关注度
        
        按以下格式输出三个选题：
        [选题1]
        标题：(选题标题)
        核心问题：(选题要解决的核心问题)
        数据新闻价值：(为什么这个选题值得做数据新闻)
        
        [选题2]
        ...
        
        [选题3]
        ...
        """
        
        messages = [
            {"role": "system", "content": system_prompt}
        ]
        
        # 添加搜索上下文（如果有）
        if search_context:
            messages.append({"role": "system", "content": search_context})
        
        messages.append({"role": "user", "content": f"我想做一个关于以下主题的数据新闻：{topic_description}"})
        
        # 显示助手正在输入的消息
        with st.chat_message("assistant"):
            suggestion_text_container = st.empty()
            suggestion_text = ""
            
            # 使用智谱API的流式输出
            for token in client.chat_completions_create(messages, stream=True):
                suggestion_text += token
                suggestion_text_container.markdown(suggestion_text)
            
            # 保存完整回复到会话状态
            st.session_state.topic_conversation.append({"role": "assistant", "content": suggestion_text})
            
            # 解析建议的选题
            topics = []
            pattern = r'\[选题(\d+)\](.*?)(?=\[选题\d+\]|$)'
            matches = re.findall(pattern, suggestion_text, re.DOTALL)
            
            for _, topic_content in matches:
                # 提取选题信息
                title_match = re.search(r'标题：(.*?)(?:\n|$)', topic_content)
                title = title_match.group(1).strip() if title_match else "未命名选题"
                topics.append(title)
            
            st.session_state.suggested_topics = topics
    
    # 如果有建议的选题，提供选择按钮
    if st.session_state.suggested_topics:
        st.subheader("请选择一个选题，或重新生成")
        
        cols = st.columns(3)
        for i, topic in enumerate(st.session_state.suggested_topics):
            with cols[i]:
                if st.button(f"选择: {topic}"):
                    st.session_state.selected_topic = topic
                    st.rerun()
        
        if st.button("重新生成选题"):
            # 清除之前的建议，保留对话历史
            st.session_state.suggested_topics = []
            st.rerun()
    
    # 如果用户还没有选择选题，返回False
    return False

# 数据收集方向生成阶段
def data_collection_phase():
    st.header("第二步，整理数据收集思路")
    
    # 初始化session state变量
    if 'data_directions' not in st.session_state:
        st.session_state.data_directions = None
    
    if 'data_collection_completed' not in st.session_state:
        st.session_state.data_collection_completed = False
        
    if 'data_conversation' not in st.session_state:
        st.session_state.data_conversation = []
    
    # 如果已经完成数据收集方向生成，显示结果并隐藏输入框
    if st.session_state.data_collection_completed:
        # 显示对话历史
        for message in st.session_state.data_conversation:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])
                
        # 提醒用户收集数据并上传
        st.info("👆 请参考上述数据收集方向指南，收集所需数据后通过左侧边栏上传，继续后续分析步骤。")
        
        # 添加刷新图标和提示，替代输入框
        refresh_col1, refresh_col2 = st.columns([1, 10])
        with refresh_col1:
            if st.button("🔄", help="重新生成数据收集方向"):
                st.session_state.data_directions = None
                st.session_state.data_collection_completed = False
                st.rerun()
        with refresh_col2:
            st.write("如需重新生成数据收集方向，请点击左侧刷新按钮")
        
        return True
    
    # 当用户已经选择了选题，但还没有生成数据收集方向
    if st.session_state.selected_topic and not st.session_state.data_directions:
        topic = st.session_state.selected_topic
        
        # 显示已有对话
        for message in st.session_state.data_conversation:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])
        
        # 系统提示：告知用户当前选题
        if len(st.session_state.data_conversation) == 0:
            with st.chat_message("assistant"):
                st.markdown(f"基于您选择的选题: **{topic}**，我可以帮您生成详细的数据收集方向。")
                st.session_state.data_conversation.append({
                    "role": "assistant", 
                    "content": f"基于您选择的选题: **{topic}**，我可以帮您生成详细的数据收集方向。"
                })
        
        # 用户输入或生成按钮
        user_input = st.chat_input("输入任何问题或点击'生成数据收集方向'按钮", key="data_input")
        generate_button = st.button("生成数据收集方向")
        
        if generate_button or user_input:
            if user_input:
                # 添加用户输入到对话
                st.session_state.data_conversation.append({"role": "user", "content": user_input})
                with st.chat_message("user"):
                    st.markdown(user_input)
                    
                # 进行普通回复
                with st.chat_message("assistant"):
                    response_container = st.empty()
                    response_text = ""
                    
                    # 构建普通对话提示
                    chat_messages = [
                        {"role": "system", "content": f"你是数据新闻专家，正在帮助用户规划选题'{topic}'的数据收集。回答用户所有关于数据收集的问题。"},
                    ]
                    
                    # 添加历史对话
                    for msg in st.session_state.data_conversation:
                        chat_messages.append({"role": msg["role"], "content": msg["content"]})
                    
                    # 使用流式API
                    for token in client.chat_completions_create(chat_messages, stream=True):
                        response_text += token
                        response_container.markdown(response_text)
                    
                    st.session_state.data_conversation.append({"role": "assistant", "content": response_text})
            else:
                # 生成数据收集方向
                # 构建提示
                system_prompt = """你是一位专业的数据新闻记者，擅长规划数据新闻报道的数据收集策略。

                首先，请判断用户选择的选题属于：
                - 📊 数据驱动型：从数据集出发，没有预设结论，通过数据探索发现故事
                - 💡 话题驱动型：基于明确的议题，收集数据来佐证或分析特定现象

                然后，根据用户选择的数据新闻选题，生成6-8个具体的数据收集方向，按照合理的故事递进顺序排列：

                对每个数据收集方向，请注明：
                - 类型：1. 🌐 二手数据：提供可能存在相关数据的报告、具体网站、数据库或开放数据平台，附上数据获取方法。2. 🔍 调研数据：明确是需要线下走访、网络内容分析还是问卷发放，并提供调研的重点问题和方法。3. 🤖 自主数据挖掘：推荐适合爬虫收集的网站，说明可以获取什么类型的数据，以及大致的技术难度。
                - 该数据将回答什么具体问题
                - 数据获取的可行性评估（易/中/难）
                - 获取此数据可能遇到的挑战
                - 数据处理建议

                以Markdown格式输出，每个类别使用三级标题，每个具体方向使用四级标题，并使用表格或列表呈现详细信息。
                """
                
                messages = [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"请为我的数据新闻选题《{topic}》提供数据收集方向建议。"}
                ]
                
                # 显示进度条
                with st.chat_message("assistant"):
                    directions_container = st.empty()
                    directions_text = ""
                    
                    # 使用流式输出
                    for token in client.chat_completions_create(messages, stream=True):
                        directions_text += token
                        directions_container.markdown(directions_text)
                    
                    # 确保在流式输出完成后设置状态
                    st.session_state.data_directions = directions_text
                    st.session_state.data_collection_completed = True
                    st.session_state.data_conversation.append({
                        "role": "assistant", 
                        "content": directions_text
                    })
                
                # 强制重新加载页面以应用新状态
                st.rerun()
    
    return st.session_state.data_collection_completed

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()
ZHIPU_API_KEY = '3a1df8f109f445f4b4eb898939a28a9f.0O5igS77SZZ0WGzV'  # 替换为您的API密钥
API_URL = 'https://open.bigmodel.cn/api/paas/v4/chat/completions'  # 智谱AI的API地址

# 修改OpenAI客户端初始化部分
class ZhipuClient:
    def __init__(self, api_key):
        self.api_key = api_key
        self.headers = {
            'Authorization': f'Bearer {api_key}',
            'Content-Type': 'application/json'
        }
        self.api_url = "https://open.bigmodel.cn/api/paas/v4/chat/completions"  # 智谱API URL

    def chat_completions_create(self, messages, model="glm-4-plus", temperature=0.7, stream=False):
        data = {
            "model": model,
            "messages": messages,
            "temperature": temperature,
            "stream": stream
        }

        try:
            if not stream:
                # 非流式响应处理
                response = requests.post(self.api_url, headers=self.headers, json=data)
                response.raise_for_status()
                return response.json()
            else:
                # 流式响应处理
                response = requests.post(self.api_url, headers=self.headers, json=data, stream=True)
                response.raise_for_status()
                
                for line in response.iter_lines():
                    if line:
                        line = line.decode('utf-8')
                        if line.startswith('data: '):
                            json_str = line[6:]  # 去掉'data: '前缀
                            if json_str.strip() == '[DONE]':
                                break
                            try:
                                json_data = json.loads(json_str)
                                if 'choices' in json_data and len(json_data['choices']) > 0:
                                    content = json_data['choices'][0].get('delta', {}).get('content', '')
                                    if content:
                                        yield content
                            except json.JSONDecodeError:
                                pass
                        
        except Exception as e:
            if stream:
                yield f"API调用失败: {str(e)}"
            else:
                raise Exception(f"API调用失败: {str(e)}")

# 替换OpenAI客户端初始化
client = ZhipuClient(api_key=ZHIPU_API_KEY)

def get_data_summary(df: pd.DataFrame) -> str:
    """生成数据集的简要描述"""
    summary = []

    # 基本信息
    summary.append(f"数据集包含 {len(df)} 行，{len(df.columns)} 列")

    # 列信息
    for col in df.columns:
        col_type = df[col].dtype
        unique_count = df[col].nunique()
        null_count = df[col].isnull().sum()

        # 对于数值列，添加基本统计信息
        if pd.api.types.is_numeric_dtype(df[col]):
            stats = df[col].describe()
            col_info = (f"列 '{col}' (类型: {col_type}): "
                       f"取值范围 {stats['min']:.2f} 到 {stats['max']:.2f}, "
                       f"平均值 {stats['mean']:.2f}, "
                       f"不同值数量 {unique_count}")
        else:
            # 对于非数值列，显示唯一值数量和示例值
            sample_values = df[col].dropna().sample(min(3, unique_count)).tolist()
            col_info = (f"列 '{col}' (类型: {col_type}): "
                       f"不同值数量 {unique_count}, "
                       f"示例值: {', '.join(map(str, sample_values))}")

        if null_count > 0:
            col_info += f", 存在 {null_count} 个空值"

        summary.append(col_info)

    return "\n".join(summary)

def format_visualization_suggestions(response_text: str) -> str:
    """将API响应格式化为HTML样式的输出"""

    # 定义CSS样式
    css = """
        <style>
            .suggestion {
                background-color: #f8f9fa;
                padding: 15px 20px;
                margin-bottom: 20px;
                border-left: 4px solid #4A90E2;
                border-radius: 4px;
            }
            .suggestion-number {
                font-size: 18px;
                font-weight: bold;
                color: #2c3e50;
                margin-bottom: 15px;
            }
            .label {
                color: #4A90E2;
                font-weight: bold;
                margin-top: 10px;
            }
            .content {
                color: #333;
                margin-bottom: 10px;
                line-height: 1.5;
            }
        </style>
    """

    # 将文本分割成不同的建议
    suggestions = response_text.split("\n\n---\n\n")

    html_parts = [css]

    for i, suggestion in enumerate(suggestions, 1):
        # 开始新的建议区块
        html_parts.append(f'<div class="suggestion">')

        # 解析每个部分
        sections = suggestion.strip().split("\n\n")

        # 先添加建议编号
        html_parts.append(f'<div class="suggestion-number">建议 {i}</div>')

        # 处理每个部分
        for section in sections:
            if "[" in section and "]" in section:
                header = section[section.find("[")+1:section.find("]")]
                content = section[section.find("]")+1:].strip()
                html_parts.append(f'<div class="label">{header}</div>')
                html_parts.append(f'<div class="content">{content}</div>')

        html_parts.append('</div>')

    return "".join(html_parts)

# 修改get_llm_response函数中的prompt部分
def get_llm_response(prompt: str, df: Optional[pd.DataFrame] = None) -> str:
    """获取LLM的可视化建议，使用流式输出"""
    try:
        # 如果提供了DataFrame，生成数据概要
        if df is not None:
            data_summary = get_data_summary(df)
            full_prompt = f"""请作为数据可视化专家分析以下数据集：

数据集概要：
{data_summary}

用户问题：
{prompt}"""
        else:
            full_prompt = prompt

        # 构建完整的消息数组
        messages = [
            {
                "role": "system",
                "content": """你是一个中国数据新闻专家。请分析数据并提供3-4个具体的数据可视化建议。

每个建议必须按照以下固定格式输出，确保每个部分都另起新行：

[标题]
(带有探索性与新闻价值的标题)

[使用列]
(明确指出使用哪些列)

[图表类型]
(推荐使用的图表类型，如折线图、柱状图、散点图等)

[缘由]
(解释为什么这个可视化方案有价值)

---

建议 2：
(按相同格式继续...)"""
            },
            {"role": "user", "content": full_prompt}
        ]

        # 使用流式输出
        visualization_text = ""
        for token in client.chat_completions_create(messages, model="glm-4-plus", stream=True):
            visualization_text += token
        
        # 格式化可视化建议
        return format_visualization_suggestions(visualization_text)
    except Exception as e:
        logger.error(f"LLM API 错误: {str(e)}")
        return None

# 修改cached_api_call函数
@functools.lru_cache(maxsize=32)
def cached_api_call(prompt: str) -> str:
    """缓存API调用结果"""
    try:
        response = get_llm_response(prompt)
        if response is not None:
            return response
        else:
            st.error("无法获取AI建议，请稍后重试")
            return "无法获取AI建议，请稍后重试"
    except Exception as e:
        st.error(f"API调用错误: {str(e)}")
        logger.error(f"API调用错误: {str(e)}")
        return "API调用出现错误，请检查API密钥配置或网络连接"

# Data processing class
class DataProcessor:
    def __init__(self, file_or_df):
        self.df = None
        if isinstance(file_or_df, pd.DataFrame):
            # 直接使用DataFrame，与文件上传保持一致的处理方式
            self.df = file_or_df
            self.clean_data()  # 使用统一的数据清理方法
        else:
            self.file_type = file_or_df.name.split('.')[-1].lower()
            self.process_file(file_or_df)

    def clean_data(self):
        """统一的数据清理方法"""
        if self.df is not None:
            # 清理列名
            self.df.columns = self.df.columns.astype(str)
            self.df.columns = [col.strip() for col in self.df.columns]

            # 对每列进行基础处理
            for col in self.df.columns:
                # 处理日期时间列
                if any(keyword in col.lower() for keyword in ['time', 'date']):
                    try:
                        self.df[col] = pd.to_datetime(self.df[col])
                    except:
                        continue

                # 尝试转换为数值类型（如果适合的话）
                elif self.df[col].dtype == 'object':
                    try:
                        numeric_values = pd.to_numeric(self.df[col], errors='coerce')
                        if numeric_values.notna().sum() / len(numeric_values) > 0.5:
                            self.df[col] = numeric_values
                    except:
                        continue

    def process_file(self, file):
        """处理上传的文件"""
        try:
            # 读取文件
            if self.file_type == 'csv':
                self.df = pd.read_csv(file, encoding='utf-8')
            elif self.file_type == 'xlsx':
                self.df = pd.read_excel(file, engine='openpyxl')
            elif self.file_type == 'xls':
                self.df = pd.read_excel(file, engine='xlrd')
            elif self.file_type == 'json':
                self.df = pd.read_json(file)
            else:
                st.error("不支持的文件类型。请上传 CSV、XLSX、XLS 或 JSON 文件。")
                return

            self.clean_data()  # 使用统一的数据清理方法

        except UnicodeDecodeError:
            try:
                if self.file_type == 'csv':
                    self.df = pd.read_csv(file, encoding='gbk')
                    self.clean_data()
            except Exception as e:
                st.error(f"文件编码错误: {str(e)}")
                logger.error(f"文件编码错误: {str(e)}")
                raise
        except Exception as e:
            st.error(f"文件处理错误: {str(e)}")
            logger.error(f"文件处理错误: {str(e)}")
            raise

    def get_data_profile(self) -> Dict:
        """Generate basic data profile"""
        if self.df is None:  # Handle case where file processing failed
            return {}

        profile = {
            'columns': list(self.df.columns),
            'dtypes': {str(k): str(v) for k, v in self.df.dtypes.to_dict().items()},
            'null_counts': self.df.isnull().sum().to_dict(),
        }

        numeric_cols = self.df.select_dtypes(include=['int64', 'float64']).columns
        if not numeric_cols.empty:
            profile['statistics'] = {
                col: {
                    str(k): float(v) if pd.notnull(v) else None
                    for k, v in self.df[col].describe().to_dict().items()
                }
                for col in numeric_cols
            }

        return profile

# Visualization Generator (improved)
class VisualizationGenerator:
    def __init__(self, df: pd.DataFrame):
        self.df = df
        self.color_schemes = {
            'nyt': ['#1f77b4', '#ff7f0e', '#2ca02c', '#d62728', '#9467bd', '#8c564b', '#e377c2', '#7f7f7f', '#bcbd22', '#17becf'],
            'modern': ['#4A90E2', '#50E3C2', '#F5A623', '#D0021B', '#9013FE', '#417505', '#7ED321', '#BD10E0', '#8B572A', '#4A4A4A'],
            'soft': ['#ADD8E6', '#FF9999', '#FFB6C1', '#98D8C8', '#B0E0E6', '#FFDAB9', '#DDA0DD', '#E6E6FA', '#F0E68C', '#E0FFFF']
        }
        self.current_theme = 'modern'  # 默认主题
        self.theme_map = {
            'modern': ThemeType.LIGHT,
            'nyt': ThemeType.DARK,
            'soft': ThemeType.ESSOS
        }

    def set_theme(self, theme_name: str):
        """设置当前主题"""
        if theme_name in self.color_schemes:
            self.current_theme = theme_name

    def analyze_column(self, column: str) -> dict:
        series = self.df[column]
        unique_count = series.nunique()
        total_count = len(series)
        is_numeric = pd.api.types.is_numeric_dtype(series)

        return {
            'unique_count': unique_count,
            'total_count': total_count,
            'is_numeric': is_numeric,
            'dtype': str(series.dtype)
        }

    def preprocess_categorical_data(self, column: str) -> pd.DataFrame:
        """处理分类数据"""
        # 直接使用value_counts()获取分类统计
        value_counts = self.df[column].value_counts()
        
        # 如果类别过多，只保留前10个
        if len(value_counts) > 10:
            value_counts = value_counts.head(10)
        
        return pd.DataFrame({
            'category': value_counts.index,
            'count': value_counts.values
        })

    def suggest_chart_type(self, columns: List[str]) -> str:
        """根据数据特征自动推荐图表类型"""
        if len(columns) == 1:
            column = columns[0]
            analysis = self.analyze_column(column)
            
            if not analysis['is_numeric']:
                if analysis['unique_count'] <= 10:  # 少量分类
                    return 'pie' if analysis['unique_count'] <= 6 else 'bar'
                else:  # 大量分类
                    return 'bar'
            else:  # 数值数据
                if analysis['unique_count'] > 10:  # 连续数值
                    return 'histogram'
                else:  # 离散数值
                    return 'bar'
        else:  # 双变量分析
            x_col, y_col = columns[:2]
            x_analysis = self.analyze_column(x_col)
            y_analysis = self.analyze_column(y_col)
            
            if x_analysis['is_numeric'] and y_analysis['is_numeric']:
                return 'scatter'
            elif (not x_analysis['is_numeric']) and y_analysis['is_numeric']:
                return 'box' if x_analysis['unique_count'] <= 10 else 'violin'
            elif x_analysis['is_numeric'] and (not y_analysis['is_numeric']):
                return 'bar'
            else:
                return 'heatmap' if (x_analysis['unique_count'] <= 20 and y_analysis['unique_count'] <= 20) else 'bar'

    def generate_chart(self, columns: list, chart_type: str = 'auto', title: str = None, source: str = None, show_legend: bool = True):
        """使用pyecharts生成图表"""
        try:
            # 如果是自动模式，推荐图表类型
            if chart_type == 'auto':
             chart_type = self.suggest_chart_type(columns)

            # 设置默认宽高
            width = "100%"
            height = "500px"
        
            # 图表标题
            title_text = title or (f'{columns[0]} 分布' if len(columns) == 1 else f'{columns[1]} vs {columns[0]}')
        
            # 选择当前主题的颜色方案
            colors = self.color_schemes[self.current_theme]
            theme = self.theme_map[self.current_theme]
        
            # 设置图例选项
            legend_opts = opts.LegendOpts(is_show=show_legend, pos_bottom="10%", orient="horizontal", pos_left="center")
            
            if len(columns) == 1:
                column = columns[0]
                analysis = self.analyze_column(column)
                
                # 检查数据类型与图表类型的匹配性
                if not analysis['is_numeric'] and chart_type in ['histogram', 'box', 'violin', 'heatmap']:
                    st.warning("提示：文本/分类数据不适合使用数值型图表（直方图、箱线图、小提琴图、热力图）进行展示。请选择饼图、柱状图等分类图表。")
                    return None
                
                if not analysis['is_numeric']:
                    processed_data = self.preprocess_categorical_data(column)
                    
                    if chart_type == 'pie':
                        # 创建饼图
                        chart = (
                            Pie(init_opts=opts.InitOpts(width=width, height=height, theme=theme))
                            .add(
                                series_name=column,
                                data_pair=[list(z) for z in zip(processed_data['category'], processed_data['count'])],
                                radius=["40%", "70%"],
                            )
                            .set_global_opts(
                                title_opts=opts.TitleOpts(title=title_text, pos_left="center"),
                                legend_opts=opts.LegendOpts(pos_bottom="10%", orient="horizontal", pos_left="center"),
                            )
                            .set_series_opts(label_opts=opts.LabelOpts(formatter="{b}: {c} ({d}%)"))
                        )
                        
                    elif chart_type == 'bar':
                        # 创建柱状图
                        chart = (
                            Bar(init_opts=opts.InitOpts(width=width, height=height, theme=theme))
                            .add_xaxis(processed_data['category'].tolist())
                            .add_yaxis(column, processed_data['count'].tolist())
                            .set_global_opts(
                                title_opts=opts.TitleOpts(title=title_text, pos_left="center"),
                                xaxis_opts=opts.AxisOpts(axislabel_opts=opts.LabelOpts(rotate=0)),
                                legend_opts=opts.LegendOpts(is_show=False),
                            )
                        )
                        
                    elif chart_type == 'treemap':
                        # 创建树图
                        data = [{"name": str(c), "value": int(v)} for c, v in zip(processed_data['category'], processed_data['count'])]
                        chart = (
                            TreeMap(init_opts=opts.InitOpts(width=width, height=height, theme=theme))
                            .add(
                                series_name=column,
                                data=data,
                                visual_min=0,
                                visual_max=max(processed_data['count']),
                                label_opts=opts.LabelOpts(position="inside"),
                            )
                            .set_global_opts(title_opts=opts.TitleOpts(title=title_text))
                        )
                    
                    elif chart_type == 'sunburst':
                        # 创建旭日图
                        data = [{"name": str(c), "value": int(v)} for c, v in zip(processed_data['category'], processed_data['count'])]
                        chart = (
                            Sunburst(init_opts=opts.InitOpts(width=width, height=height, theme=theme))
                            .add(
                                series_name=column,
                                data_pair=data,
                                radius=[0, "90%"],
                            )
                            .set_global_opts(title_opts=opts.TitleOpts(title=title_text))
                        )
                
                elif analysis['is_numeric']:
                    if chart_type == 'histogram':
                        # 处理直方图，使用Bar实现
                        # 生成直方图数据
                        hist, bin_edges = np.histogram(self.df[column].dropna(), bins='auto')
                        bin_labels = [f"{bin_edges[i]:.2f}-{bin_edges[i+1]:.2f}" for i in range(len(bin_edges)-1)]
                        
                        chart = (
                            Bar(init_opts=opts.InitOpts(width=width, height=height, theme=theme))
                            .add_xaxis(bin_labels)
                            .add_yaxis("频率", hist.tolist())
                            .set_global_opts(
                                title_opts=opts.TitleOpts(title=title_text, pos_left="center"),
                                xaxis_opts=opts.AxisOpts(axislabel_opts=opts.LabelOpts(rotate=45)),
                            )
                        )
                    
                    elif chart_type == 'box':
                        # 箱线图数据准备
                        data = self.df[column].dropna().tolist()
                        chart = (
                            Boxplot(init_opts=opts.InitOpts(width=width, height=height, theme=theme))
                            .add_xaxis([column])
                            .add_yaxis("", self._prepare_boxplot_data(data))
                            .set_global_opts(
                                title_opts=opts.TitleOpts(title=title_text, pos_left="center"),
                                yaxis_opts=opts.AxisOpts(name=column),
                            )
                        )
                    
                    elif chart_type == 'violin':
                        # pyecharts不直接支持小提琴图，这里用boxplot替代
                        data = self.df[column].dropna().tolist()
                        st.warning("ECharts不直接支持小提琴图，已替换为箱线图展示")
                        chart = (
                            Boxplot(init_opts=opts.InitOpts(width=width, height=height, theme=theme))
                            .add_xaxis([column])
                            .add_yaxis("", self._prepare_boxplot_data(data))
                            .set_global_opts(
                                title_opts=opts.TitleOpts(title=title_text, pos_left="center"),
                                yaxis_opts=opts.AxisOpts(name=column),
                            )
                        )
                
            else:  # 双变量分析
                x_col, y_col = columns[:2]
                
                if chart_type == 'scatter':
                    # 散点图
                    chart = (
                        Scatter(init_opts=opts.InitOpts(width=width, height=height, theme=theme))
                        .add_xaxis(self.df[x_col].tolist())
                        .add_yaxis(
                            y_col,
                            self.df[[x_col, y_col]].dropna().values.tolist(),
                            symbol_size=10,
                        )
                        .set_global_opts(
                            title_opts=opts.TitleOpts(title=title_text, pos_left="center"),
                            xaxis_opts=opts.AxisOpts(name=x_col),
                            yaxis_opts=opts.AxisOpts(name=y_col),
                            visualmap_opts=opts.VisualMapOpts(type_="size", max_=100, min_=10),
                        )
                    )
                
                elif chart_type == 'line':
                    # 折线图
                    chart = (
                        Line(init_opts=opts.InitOpts(width=width, height=height, theme=theme))
                        .add_xaxis(self.df[x_col].tolist())
                        .add_yaxis(
                            y_col,
                            self.df[y_col].tolist(),
                            symbol_size=8,
                        )
                        .set_global_opts(
                            title_opts=opts.TitleOpts(title=title_text, pos_left="center"),
                            xaxis_opts=opts.AxisOpts(
                                name=x_col,
                                type_="category" if not pd.api.types.is_numeric_dtype(self.df[x_col]) else "value"
                            ),
                            yaxis_opts=opts.AxisOpts(name=y_col),
                            datazoom_opts=[opts.DataZoomOpts()],
                            legend_opts=opts.LegendOpts(pos_bottom="10%", orient="horizontal", pos_left="center"),
                        )
                    )
                
                elif chart_type == 'bar':
                    # 柱状图
                    chart = (
                        Bar(init_opts=opts.InitOpts(width=width, height=height, theme=theme))
                        .add_xaxis(self.df[x_col].astype(str).tolist())
                        .add_yaxis(y_col, self.df[y_col].tolist())
                        .set_global_opts(
                            title_opts=opts.TitleOpts(title=title_text, pos_left="center"),
                            xaxis_opts=opts.AxisOpts(
                                name=x_col,
                                axislabel_opts=opts.LabelOpts(rotate=45)
                            ),
                            yaxis_opts=opts.AxisOpts(name=y_col),
                            datazoom_opts=[opts.DataZoomOpts()],
                            legend_opts=opts.LegendOpts(pos_bottom="10%", orient="horizontal", pos_left="center"),
                        )
                    )
                
                elif chart_type == 'box':
                    # 分组箱线图 - 简化实现
                    st.warning("ECharts中的分组箱线图实现较为复杂，展示效果可能与预期有差异")
                    grouped = self.df.groupby(x_col)[y_col].apply(list).reset_index()
                    chart = (
                        Boxplot(init_opts=opts.InitOpts(width=width, height=height, theme=theme))
                        .add_xaxis(grouped[x_col].tolist())
                        .add_yaxis(
                            y_col,
                            [self._prepare_boxplot_data(data)[0] for data in grouped[y_col]]
                        )
                        .set_global_opts(
                            title_opts=opts.TitleOpts(title=title_text, pos_left="center"),
                            xaxis_opts=opts.AxisOpts(name=x_col),
                            yaxis_opts=opts.AxisOpts(name=y_col),
                        )
                    )
                
                elif chart_type == 'violin':
                    # 简化实现：pyecharts不直接支持小提琴图
                    st.warning("ECharts不直接支持小提琴图，已替换为分组箱线图展示")
                    grouped = self.df.groupby(x_col)[y_col].apply(list).reset_index()
                    chart = (
                        Boxplot(init_opts=opts.InitOpts(width=width, height=height, theme=theme))
                        .add_xaxis(grouped[x_col].tolist())
                        .add_yaxis(
                            y_col,
                            [self._prepare_boxplot_data(data)[0] for data in grouped[y_col]]
                        )
                        .set_global_opts(
                            title_opts=opts.TitleOpts(title=title_text, pos_left="center"),
                            xaxis_opts=opts.AxisOpts(name=x_col),
                            yaxis_opts=opts.AxisOpts(name=y_col),
                        )
                    )
                
                elif chart_type == 'heatmap':
                    # 热力图
                    # 简化实现：创建数据透视表
                    if pd.api.types.is_numeric_dtype(self.df[y_col]):
                        # 如果y是数值列，计算平均值
                        pivot_data = self.df.pivot_table(
                            values=y_col,
                            index=x_col,
                            aggfunc='mean'
                        ).reset_index()
                        x_data = pivot_data[x_col].astype(str).tolist()
                        y_data = [y_col]
                        heat_data = [[0, 0, val] for val in pivot_data[y_col]]
                    else:
                        # 如果y是分类列，计算频数
                        counts = self.df.groupby([x_col, y_col]).size().reset_index(name='count')
                        x_data = sorted(counts[x_col].unique().astype(str).tolist())
                        y_data = sorted(counts[y_col].unique().astype(str).tolist())
                        heat_data = []
                        for _, row in counts.iterrows():
                            x_idx = x_data.index(str(row[x_col]))
                            y_idx = y_data.index(str(row[y_col]))
                            heat_data.append([x_idx, y_idx, row['count']])
                    
                    chart = (
                        HeatMap(init_opts=opts.InitOpts(width=width, height=height, theme=theme))
                        .add_xaxis(x_data)
                        .add_yaxis(
                            "",
                            y_data,
                            heat_data,
                        )
                        .set_global_opts(
                            title_opts=opts.TitleOpts(title=title_text, pos_left="center"),
                            visualmap_opts=opts.VisualMapOpts(),
                            xaxis_opts=opts.AxisOpts(name=x_col),
                            yaxis_opts=opts.AxisOpts(name=y_col),
                        )
                    )

            # 添加数据来源注释
            if source:
                chart.set_global_opts(
                    title_opts=opts.TitleOpts(
                        title=title_text,
                        subtitle=f"数据来源: {source}",
                        pos_left="center",
                        title_textstyle_opts=opts.TextStyleOpts(font_size=22)  # 设置更大的标题字号
                    )
                )

            return chart

        except Exception as e:
            st.error(f"生成图表时出错: {str(e)}")
            logger.error(f"图表生成错误: {str(e)}")
            return None

    # 辅助方法：为箱线图准备数据
    def _prepare_boxplot_data(self, data):
        """为ECharts箱线图准备数据"""
        result = []
        if not data:
            return [[0, 0, 0, 0, 0]]
            
        data = sorted(data)
        q1, q2, q3 = np.percentile(data, [25, 50, 75])
        iqr = q3 - q1
        low_whisker = max(min(data), q1 - 1.5 * iqr)
        high_whisker = min(max(data), q3 + 1.5 * iqr)
        result.append([low_whisker, q1, q2, q3, high_whisker])
        return result

    def evaluate_chart(self, chart_type: str, columns: List[str]) -> Tuple[str, List[str], str]:
        """评估图表适用性"""
        if chart_type == 'auto':
            return "非常适合", ["数据分析价值", "图表类型适用性"], "自动选择的图表类型最适合当前数据特征。"

        data_types = [self.df[col].dtype for col in columns]
        num_numeric = sum(1 for dtype in data_types if pd.api.types.is_numeric_dtype(dtype))
        num_categorical = sum(1 for dtype in data_types if dtype == 'object' or str(dtype).startswith('datetime'))
        num_columns = len(columns)
        score = "基本适合"  # 默认为"基本适合"
        feedback_dimensions = ["数据分析价值", "图表类型适用性"]
        feedback = ""

        if chart_type == 'line':
            # 检查X轴是否为时间类型
            if num_columns == 2:
                x_col = columns[0]
                is_time_col = pd.api.types.is_datetime64_any_dtype(self.df[x_col]) or \
                             any(keyword in x_col.lower() for keyword in ['time', 'date', '时间', '日期', '年', '月'])
                if not is_time_col:
                    score = "不适合"
                    feedback = "折线图最适合展示随时间变化的趋势。当前X轴不是时间类型，建议使用其他图表类型。"
                else:
                    score = "非常适合"
                    feedback = "折线图很好地展示了数据随时间的变化趋势。"
            else:
                score = "不适合"
                feedback = "折线图需要一个时间类型的X轴和一个数值类型的Y轴。"
        elif chart_type == 'pie':
            if num_columns != 1 or num_categorical != 1:
                score = "不适合"
                feedback = "饼图最适合展示单个分类变量的分布情况。"
            else:
                score = "非常适合"
                feedback = "饼图完美展现了单个分类变量的分布比例。"
        elif chart_type == 'bar':
            if num_columns == 1 and num_categorical == 1:
                score = "非常适合"
                feedback = "柱状图很好地展示了分类数据的对比。"
            elif num_columns == 2 and num_categorical == 1 and num_numeric == 1:
                score = "非常适合"
                feedback = "柱状图有效地展示了不同类别的数值对比。"
            else:
                score = "基本适合"
                feedback = "柱状图可以展示当前数据，但可能存在更好的可视化方式。"
        elif chart_type == 'scatter':
            if num_columns == 2 and num_numeric == 2:
                score = "非常适合"
                feedback = "散点图完美展示了两个数值变量间的关系。"
            else:
                score = "不适合"
                feedback = "散点图仅适用于展示两个数值变量的关系。"
        elif chart_type == 'histogram':
            if num_columns == 1 and num_numeric == 1:
                score = "非常适合"
                feedback = "直方图很好地展示了数值变量的分布情况。"
            else:
                score = "不适合"
                feedback = "直方图仅适用于展示单个数值变量的分布。"
        else:  # 其他图表类型
            score = "基本适合"
            feedback = "此图表类型可以展示当前数据。"

        return score, feedback_dimensions, feedback

    def get_chart_data(self, columns: List[str]) -> pd.DataFrame:
        """获取用于生成数据故事的图表数据"""
        if len(columns) == 1:
            column = columns[0]
            if not pd.api.types.is_numeric_dtype(self.df[column]):
                # 对于分类数据，计算频率和百分比
                value_counts = self.df[column].value_counts()
                percentages = value_counts / len(self.df) * 100
                return pd.DataFrame({
                    'category': value_counts.index,
                    'count': value_counts.values,
                    'percentage': percentages.values
                })
            else:
                # 对于数值数据，计算基本统计量
                stats = self.df[column].describe()
                return pd.DataFrame({
                    '统计指标': stats.index,
                    '值': stats.values
                })
        else:
            # 对于双变量分析，返回原始数据的相关部分
            return self.df[columns].copy()

def simulate_progress_bar():
    """模拟进度条动画"""
    progress_bar = st.progress(0)
    progress_text = st.empty()
    progress = 0

    while progress < 90:
        # 非线性进度增加，开始快，后面慢
        increment = max(0.3, (90 - progress) / 50)
        progress = min(90, progress + increment)

        # 更新进度条和文本
        progress_bar.progress(int(progress))
        progress_text.text(f'分析进度：{int(progress)}%')
        time.sleep(0.2)

    return progress_bar, progress_text

def get_data_story(chart_config: dict, data: pd.DataFrame, evaluation_score: str) -> str:
    """生成数据故事"""
    try:
        if data.empty:
            st.warning("没有足够的数据来生成故事。")
            return None

        # 构建数据概要字符串
        data_summary = "数据分析结果：\n"
        
        # 根据数据类型构建不同的描述
        if 'percentage' in data.columns:
            # 分类数据的描述
            total_count = data['count'].sum()
            data_summary += f"总计样本数：{total_count}\n\n"
            data_summary += "类别分布：\n"
            for _, row in data.iterrows():
                data_summary += f"- {row['category']}: {row['count']}次 (占比{row['percentage']}%)\n"
        
        elif '统计指标' in data.columns:
            # 数值数据的描述
            data_summary += "数值统计：\n"
            for _, row in data.iterrows():
                data_summary += f"- {row['统计指标']}: {row['值']}\n"
        
        else:
            # 其他数据类型的描述
            data_summary += data.to_string()

        # 构建Prompt
        prompt = f"""作为一名专业的数据新闻记者，请基于以下图表信息撰写一段数据新闻段落。
图表信息：
- 标题：{chart_config.get('title', '')}
- 图表类型：{chart_config.get('chart_type', '')}
- 使用数据列：{', '.join(chart_config.get('columns', []))}
- 数据来源：{chart_config.get('source', '')}

{data_summary}

要求：
1. 使用中文数据新闻专业写作风格
2. 突出数据发现的新闻价值
3. 客观陈述，准确引用数据
4. 注重数据背后的故事性
5. 语言简洁专业"""

        logger.info(f"发送给智谱AI的Prompt: \n{prompt}")

        # 调用智谱AI
        response = client.chat_completions_create(
            messages=[
                {"role": "system", "content": "你是一位经验丰富的数据新闻记者，擅长将数据分析转化为引人入胜的新闻故事。"},
                {"role": "user", "content": prompt}
            ],
            model="glm-4-plus",
            temperature=0.7
        )

        if 'choices' in response and len(response['choices']) > 0:
            story = response['choices'][0]['message']['content']
            logger.info(f"获得的故事内容: {story}")
            return story
        else:
            logger.error("API响应格式错误")
            st.error("生成故事时发生错误，API响应格式不正确。")
            return None

    except Exception as e:
        logger.error(f"生成故事时出错: {str(e)}")
        st.error(f"生成数据故事时发生错误: {str(e)}")
        return None

def get_data_news_story(selected_charts):
    """基于多个选定图表生成完整的数据新闻故事"""
    if not selected_charts:
        return None
    
    # 构建数据概要字符串
    charts_info = []
    
    for i, chart_info in enumerate(selected_charts):
        config = chart_info['config']
        data = chart_info['data']
        
        chart_summary = f"图表{i+1}信息：\n"
        chart_summary += f"- 标题：{config.get('title', f'图表{i+1}')}\n"
        chart_summary += f"- 图表类型：{config.get('chart_type', '自动')}\n"
        chart_summary += f"- 使用数据列：{', '.join(config.get('columns', []))}\n"
        chart_summary += f"- 数据来源：{config.get('source', '未指定')}\n\n"
        
        # 添加数据描述
        chart_summary += "数据分析结果：\n"
        
        if not data.empty:
            # 根据数据类型构建不同的描述
            if 'percentage' in data.columns:
                # 分类数据的描述
                total_count = data['count'].sum()
                chart_summary += f"总计样本数：{total_count}\n"
                chart_summary += "类别分布：\n"
                for _, row in data.head(5).iterrows():
                    chart_summary += f"- {row['category']}: {row['count']}次 (占比{row['percentage']:.2f}%)\n"
                if len(data) > 5:
                    chart_summary += f"- ... 等{len(data)}个类别\n"
            
            elif '统计指标' in data.columns:
                # 数值数据的描述
                chart_summary += "数值统计：\n"
                for _, row in data.iterrows():
                    chart_summary += f"- {row['统计指标']}: {row['值']}\n"
            
            else:
                # 其他数据类型的描述
                chart_summary += "数据预览：\n"
                chart_summary += data.head(3).to_string() + "\n...(数据省略)\n"
        
        charts_info.append(chart_summary)
    
    # 构建优化后的Prompt
    prompt = f"""作为一名专业数据新闻记者，请基于以下{len(charts_info)}个图表信息撰写一篇完整的数据新闻文章。

{"\n\n".join(charts_info)}

需求：
1. 文章需要一个吸引人的标题，使用"# 标题"格式
2. 将文章分成2-4个小节，每个小节标题使用"### 小节标题"格式
3. 文章篇幅适中（400-1100字），字数与图表数量成正比
4. 文章风格：
   - 开头引出核心发现，设置新闻基调
   - 中间部分深入分析每个图表数据，揭示数据背后的故事和关联
   - 结尾提供总结性观点或建议
5. 表达要求：
   - 客观准确引用数据，避免过度推测
   - 使用专业但通俗易懂的语言
   - 适当运用比喻、对比等修辞手法增强可读性

请直接输出完整的新闻文章，无需解释你的写作过程。"""

    logger.info(f"发送给智谱AI的Prompt: \n{prompt}")

    # 调用智谱AI
    try:
        response = client.chat_completions_create(
            messages=[
                {"role": "system", "content": "你是一位经验丰富的数据新闻记者，擅长将多维度数据分析转化为引人入胜的新闻故事。你会分析多个图表之间的关联，提炼出数据背后的深层含义。"},
                {"role": "user", "content": prompt}
            ],
            model="glm-4-plus",
            temperature=0.7
        )

        if 'choices' in response and len(response['choices']) > 0:
            story = response['choices'][0]['message']['content']
            logger.info(f"获得的新闻故事内容: {story}")
            return story
        else:
            logger.error("API响应格式错误")
            return None
    except Exception as e:
        logger.error(f"生成数据新闻时出错: {str(e)}")
        return None

# 新增爬虫数据处理类
class WebDataCrawler:
    def __init__(self):
        self.headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
    
    def crawl_data(self, url: str) -> pd.DataFrame:
        """爬取网页数据并转换为DataFrame"""
        try:
            st.write("开始爬取数据...")
            progress_bar = st.progress(0)
            
            # 发送请求获取页面内容
            response = requests.get(url, headers=self.headers, timeout=30)
            response.raise_for_status()
            
            progress_bar.progress(0.5)
            
            # 尝试提取表格数据
            tables = pd.read_html(response.text)
            if tables:
                df = tables[0]  # 获取第一个表格
                # 保存到session state
                st.session_state['crawled_df'] = df
                st.write("数据预览：")
                st.write(df.head())
                progress_bar.progress(1.0)
                return df
            
            # 如果没有表格，提取文本内容
            soup = BeautifulSoup(response.text, 'html.parser')
            text_content = soup.get_text()
            df = pd.DataFrame({
                'content': [text_content],
                'url': [url],
                'timestamp': [pd.Timestamp.now()]
            })
            
            # 保存到session state
            st.session_state['crawled_df'] = df
            progress_bar.progress(1.0)
            st.write("获取到的文本数据预览：")
            st.write(text_content[:500] + "...")
            
            return df
            
        except Exception as e:
            st.error(f"爬取失败: {str(e)}")
            logger.error(f"爬取失败: {str(e)}")
            return pd.DataFrame()

# Main Application
def main():
    # 初始化session state变量
    if 'selected_charts' not in st.session_state:
        st.session_state.selected_charts = []
    
    st.title("复新Vis-数据新闻多智能体工作流")
    
    # 使用纯白色背景，只保留介绍文字的样式
    st.markdown(
        """
        <style>
        .intro-text {
            padding: 20px;
            border-radius: 10px;
            background-color: rgba(255, 255, 255, 0.9);
            margin: 20px 0;
        }
        
        .intro-point {
            margin: 10px 0;
            padding-left: 20px;
        }
        </style>
        """,
        unsafe_allow_html=True
    )

    # 始终显示侧边栏
    with st.sidebar:
        st.header("当你准备好了，你可以开始数据输入")
        data_input_method = st.radio(
            "选择数据输入方式",
            ["上传文件", "网页爬取"]
        )
        
        if data_input_method == "上传文件":
            uploaded_file = st.file_uploader("上传 CSV、Excel 或 JSON 文件", 
                                           type=['csv', 'xlsx', 'xls', 'json'])
            if uploaded_file:
                processor = DataProcessor(uploaded_file)
                # 保存到session state
                st.session_state['current_processor'] = processor
                # 设置状态，表示已上传数据
                st.session_state['data_uploaded'] = True
                
        else:  # 网页爬取
            url = st.text_input("输入要爬取的网页URL")
            
            with st.expander("爬取配置"):
                timeout = st.slider("超时时间(秒)", 10, 60, 30)
            
            if st.button("开始爬取", key="crawl_button"):
                if url:
                    try:
                        crawler = WebDataCrawler()
                        df = crawler.crawl_data(url)
                        
                        if not df.empty:
                            processor = DataProcessor(df)
                            # 保存到session state
                            st.session_state['current_processor'] = processor
                            # 设置状态，表示已上传数据
                            st.session_state['data_uploaded'] = True
                            st.success("数据爬取成功！")
                        else:
                            st.warning("未获取到数据，请检查URL或尝试其他网页。")
                    except Exception as e:
                        st.error(f"爬取失败: {str(e)}")

    # 从session state获取processor
    processor = st.session_state.get('current_processor', None)
    
    # 如果数据未上传且未跳过选题阶段，则显示选题和数据收集界面
    if not st.session_state.get('data_uploaded', False):
        # 第一阶段：选题确定
        if not topic_selection_phase():
            return  # 如果还没完成选题确定，不进入下一阶段
        
        # 第二阶段：数据收集方向
        if not data_collection_phase():
            return  # 如果还没完成数据收集方向生成，不进入下一阶段
            
        # 如果仍然没有数据，显示介绍内容
        if not processor or processor.df is None:
            st.markdown(
                """
                <div class="intro-text">
                    <h3>接下来的工作流程（版本0226）</h3>
                    <div class="intro-point">📊 <b>首先，上传你的数据：</b>支持上传本地数据集或使用我们的网页数据爬取</div>
                    <div class="intro-point">🤖 <b>然后，获取可视化的建议：</b>大模型会基于数据特征提供专业的可视化建议</div>
                    <div class="intro-point">📈 <b>其次，制作可视化图表：</b>提供多种图表类型和漂亮的配色选择</div>
                    <div class="intro-point">📝 <b>最后，拿上评估合格的图表，撰写出数据故事：</b>自动生成专业媒体风格的数据新闻段落</div>
                </div>
                """,
                unsafe_allow_html=True
            )
            return
    
    # 如果上传了数据，显示数据处理和可视化界面
    if processor and processor.df is not None:
        # 如果之前完成了选题阶段，显示选题信息
        if st.session_state.get('selected_topic') and not st.session_state.get('skip_topic_selection'):
            with st.expander("已选定的选题", expanded=False):
                st.success(f"数据新闻选题：{st.session_state.selected_topic}")
                if st.session_state.get('data_directions'):
                    st.markdown(st.session_state.data_directions)
        
        # 1. 数据预览部分 - 默认展开
        with st.expander("数据预览", expanded=True):
            st.dataframe(processor.df.head(31), use_container_width=True)

        # 2. 数据可视化建议部分
        st.subheader("第三步，获取数据可视化建议")
        suggestion_container = st.container()
        with suggestion_container:
            if st.button("获取可视化建议", key="viz_suggestion_btn"):
                progress_bar, progress_text = simulate_progress_bar()

                response = get_llm_response("请为这个数据集提供可视化建议", processor.df)

                # 完成时将进度设为100%
                progress_bar.progress(100)
                progress_text.text('分析完成！')
                time.sleep(0.5)  # 短暂显示完成状态

                progress_bar.empty()
                progress_text.empty()

                if response:
                    st.session_state.visualization_suggestions = response
                    st.markdown(response, unsafe_allow_html=True)

            elif st.session_state.get('visualization_suggestions'):
                st.markdown(st.session_state.visualization_suggestions, unsafe_allow_html=True)

        # 3. 可视化制作部分
        st.subheader("第四步，创建可视化")
        col1, col2 = st.columns([1, 2])

        with col1:
            # 添加主题选择
            color_theme = st.selectbox(
                "选择配色主题",
                options=['modern', 'nyt', 'soft'],
                format_func=lambda x: {
                    'modern': '现代简约',
                    'nyt': '新闻专业',
                    'soft': '柔和清新'
                }[x]
            )

            show_legend = st.checkbox("显示图例", value=True)

            viz_type = st.radio(
                "选择分析类型",
                options=['单列分析', '双列关系分析'],
                horizontal=True
            )

            custom_title = st.text_input("输入图表标题（可选）", "")
            data_source = st.text_input("输入数据来源（可选）", "")

            if viz_type == '单列分析':
                column = st.selectbox("选择要分析的列", options=processor.df.columns)
                chart_type = st.radio(
                    "选择图表类型",
                    options=['自动', '饼图', '柱状图', '直方图', '箱线图', '小提琴图', '树图', '旭日图'],
                    horizontal=True
                )
                columns_to_use = [column]
            else:
                x_column = st.selectbox("选择 X 轴数据", options=processor.df.columns)
                y_column = st.selectbox("选择 Y 轴数据", options=processor.df.columns)
                chart_type = st.radio(
                    "选择图表类型",
                    options=['自动', '折线图', '柱状图', '散点图', '箱线图', '小提琴图', '热力图'],
                    horizontal=True
                )
                columns_to_use = [x_column, y_column]

            if st.button("生成图表"):
                # 每次生成新图表时，清除之前的数据故事
                if 'data_story' in st.session_state:
                    del st.session_state['data_story']
                
                st.session_state.show_legend = show_legend

                st.session_state.current_chart_config = {
                    'viz_type': viz_type,
                    'columns': columns_to_use,
                    'chart_type': chart_type,
                    'title': custom_title,
                    'source': data_source
                }

        with col2:
            if 'current_chart_config' in st.session_state:
                config = st.session_state.current_chart_config
                vis_gen = VisualizationGenerator(processor.df)
                vis_gen.set_theme(color_theme)  # 设置选择的主题

                # 转换英文图表类型为中文
                chart_type_map = {
                    '自动': 'auto',
                    '饼图': 'pie',
                    '柱状图': 'bar',
                    '直方图': 'histogram',
                    '折线图': 'line',
                    '散点图': 'scatter',
                    '箱线图': 'box',
                    '小提琴图': 'violin',
                    '树图': 'treemap',
                    '旭日图': 'sunburst'
                }

                chart_type = chart_type_map.get(config['chart_type'], config['chart_type'])

                chart = vis_gen.generate_chart(
                    columns=config['columns'],
                    chart_type=chart_type,
                    title=config['title'],
                    source=config['source'],
                    show_legend=st.session_state.get('show_legend', True)  # 获取图例显示状态
                )

                if chart:
                    # 修改：使用st_pyecharts显示ECharts图表，而不是st.plotly_chart
                    st_pyecharts(chart, height="500px")

                    with st.expander("图表评估结果（对图表点击右键可保存为图片）", expanded=True):
                        score, dimensions, feedback = vis_gen.evaluate_chart(
                            chart_type,
                            config['columns']
                        )
                        st.write(f"**图表评估得分:** {score}")
                        st.write("**评估维度:**")
                        for dim in dimensions:
                            st.write(f"- {dim}")
                        st.write(f"**评估建议:** {feedback}")

                        # 移除原有的故事生成按钮，替换为选定按钮
                        if score in ["基本适合", "非常适合"]:
                            if st.button("选定此图表"):
                                # 检查是否已经选定了5个图表
                                if len(st.session_state.selected_charts) >= 5:
                                    st.warning("最多只能选定5个图表！请先删除一些已选定的图表。")
                                else:
                                    # 将当前图表配置和数据添加到已选定图表列表
                                    chart_data = vis_gen.get_chart_data(config['columns'])
                                    # 保存图表配置、数据和评估信息
                                    chart_info = {
                                        'config': config.copy(),
                                        'data': chart_data,
                                        'score': score,
                                        'chart': chart  # 保存图表对象
                                    }
                                    st.session_state.selected_charts.append(chart_info)
                                    st.success(f"已选定图表，当前已选定 {len(st.session_state.selected_charts)} 个图表")

    # 显示已选定的图表和第三步生成数据新闻
    if processor and processor.df is not None:  # 确保有数据被加载后才执行
        if 'selected_charts' in st.session_state and st.session_state.selected_charts:
            st.subheader("已选定的图表")
            # 使用列表容器展示已选图表
            for i, chart_info in enumerate(st.session_state.selected_charts):
                col1, col2 = st.columns([3, 1])
                
                with col1:
                    st.write(f"**图表 {i+1}**: {chart_info['config'].get('title', '未命名图表')}")
                    # 显示图表
                    st_pyecharts(chart_info['chart'], height="300px")
                
                with col2:
                    st.write(f"评估: {chart_info['score']}")
                    # 添加删除按钮
                    if st.button(f"删除此图表", key=f"del_chart_{i}"):
                        st.session_state.selected_charts.pop(i)
                        st.rerun()
            
            # 第三步 - 生成完整数据新闻
            st.subheader("第五步，写作数据故事")
            
            if st.button("生成完整数据新闻"):
                progress_bar = st.progress(0)
                progress_text = st.empty()
                
                # 模拟进度
                for i in range(0, 101, 10):
                    progress_bar.progress(i)
                    if i == 0:
                        progress_text.text("初始化数据分析...")
                    elif i == 20:
                        progress_text.text("提取数据关键点...")
                    elif i == 40:
                        progress_text.text("构建新闻故事架构...")
                    elif i == 60:
                        progress_text.text("生成新闻内容...")
                    elif i == 80:
                        progress_text.text("润色文章表达...")
                    time.sleep(1.7)
                
                # 实际生成数据新闻
                story = get_data_news_story(st.session_state.selected_charts)
                
                # 完成进度
                progress_bar.progress(100)
                progress_text.text("数据新闻生成完成！")
                time.sleep(0.5)
                
                # 清除进度条和文本
                progress_bar.empty()
                progress_text.empty()
                
                if story:
                    st.session_state.news_story = story
                else:
                    st.error("无法生成数据新闻，请稍后重试。")
            
            # 显示数据新闻
            if 'news_story' in st.session_state:
                # 设计一个富媒体框来展示新闻内容
                st.markdown(
                    """
                    <style>
                    .news-container {
                        padding: 20px;
                        background-color: #f8f9fa;
                        border-radius: 10px;
                        border-left: 5px solid #4A90E2;
                        margin: 10px 0;
                    }
                    .news-title {
                        font-size: 24px;
                        font-weight: bold;
                        margin-bottom: 15px;
                        color: #2c3e50;
                    }
                    .news-section {
                        font-size: 18px;
                        font-weight: bold;
                        margin: 15px 0 10px 0;
                        color: #3498db;
                    }
                    .news-content {
                        font-size: 16px;
                        line-height: 1.6;
                        color: #333;
                    }
                    </style>
                    """,
                    unsafe_allow_html=True
                )
                
                # 处理Markdown格式的新闻内容
                news_content = st.session_state.news_story
                
                # 使用正则表达式提取标题和小节标题
                # 假设最大的标题使用# 或## 开始，小节标题使用### 开始
                title_match = re.search(r'^#\s+(.+)$|^##\s+(.+)$', news_content, re.MULTILINE)
                if title_match:
                    title = title_match.group(1) if title_match.group(1) else title_match.group(2)
                    # 从内容中移除主标题
                    news_content = re.sub(r'^#\s+.+$|^##\s+.+$', '', news_content, count=1, flags=re.MULTILINE)
                else:
                    title = "数据新闻报道"
                
                # 查找所有小节标题和内容
                sections = re.split(r'^###\s+(.+)$', news_content, flags=re.MULTILINE)
                
                # 显示富媒体格式的新闻
                news_html = f'<div class="news-container"><div class="news-title">{title}</div>'
                
                if len(sections) > 1:  # 有小节标题
                    for i in range(1, len(sections), 2):
                        if i < len(sections):
                            section_title = sections[i]
                            section_content = sections[i + 1] if i + 1 < len(sections) else ""
                            news_html += f'<div class="news-section">{section_title}</div>'
                            news_html += f'<div class="news-content">{section_content}</div>'
                else:  # 没有小节标题，直接显示内容
                    news_html += f'<div class="news-content">{news_content}</div>'
                
                news_html += '</div>'
                st.markdown(news_html, unsafe_allow_html=True)
                
                # 提供下载按钮
                col1, col2 = st.columns(2)
                with col1:
                    st.download_button(
                        label="下载Markdown格式",
                        data=st.session_state.news_story,
                        file_name="data_news_story.md",
                        mime="text/markdown"
                    )
                with col2:
                    word_file = export_to_word(st.session_state.news_story)
                    st.download_button(
                        label="下载Word文档",
                        data=word_file,
                        file_name="data_news_story.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
        elif processor and processor.df is not None:  # 只在有数据但没有选定图表时显示提示
            st.info("请先选定至少一个图表，才能生成数据新闻。")

def export_to_word(news_content, selected_charts=None):
    """将新闻内容导出为Word文档"""
    doc = Document()
    
    # 处理Markdown内容
    # 提取主标题
    title_match = re.search(r'^#\s+(.+)$|^##\s+(.+)$', news_content, re.MULTILINE)
    if title_match:
        title = title_match.group(1) if title_match.group(1) else title_match.group(2)
        doc.add_heading(title, level=0)
        # 从内容中移除主标题
        news_content = re.sub(r'^#\s+.+$|^##\s+.+$', '', news_content, count=1, flags=re.MULTILINE)
    
    # 处理小节标题和内容
    sections = re.split(r'^###\s+(.+)$', news_content, flags=re.MULTILINE)
    
    # 如果有小节
    if len(sections) > 1:
        # 处理第一个非标题部分(如果有的话)
        if sections[0].strip():
            doc.add_paragraph(sections[0].strip())
            
        # 处理各小节
        for i in range(1, len(sections), 2):
            if i < len(sections):
                section_title = sections[i]
                section_content = sections[i + 1] if i + 1 < len(sections) else ""
                
                # 添加小节标题
                doc.add_heading(section_title, level=2)
                
                # 添加小节内容
                doc.add_paragraph(section_content.strip())
    else:
        # 没有小节，直接添加内容
        doc.add_paragraph(news_content.strip())
    
    # 保存Word文档到内存
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io

if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        st.error(f"Application error: {str(e)}")
        logger.error(f"Application error: {str(e)}")